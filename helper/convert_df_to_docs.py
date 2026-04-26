from io import BytesIO

from docx import Document

def create_docx(dataframe):
    doc = Document()
    doc.add_heading("FMCG Deal Newletter", 0)

    for _, row in dataframe.iterrows():
        doc.add_heading(row['title'], level=2)
        doc.add_paragraph(f"Summary: {row['summary']}")
        doc.add_paragraph(f"source:{row['url']}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
